"""
Word output engine for DocTura Desktop.

Supports orientation control and table formatting.
"""

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table as DocxTable

from docutura.core.models import DocumentMetadata, ExtractedTable, WordOrientation
from docutura.core.themes import Theme


class WordWriter:
    """Writes extracted tables to Word documents."""

    def __init__(
        self,
        orientation: WordOrientation = WordOrientation.PORTRAIT,
        page_break_per_table: bool = False,
        include_images: bool = True,
        theme: Optional[Theme] = None,
    ):
        """
        Initialize Word writer.

        Args:
            orientation: Page orientation
            page_break_per_table: Insert page break after each table
            include_images: Include extracted images (if any)
            theme: Theme for styling
        """
        self.orientation = orientation
        self.page_break_per_table = page_break_per_table
        self.include_images = include_images
        self.theme = theme

    def write_to_word(
        self,
        output_path: Path,
        tables: List[ExtractedTable],
        metadata: Optional[DocumentMetadata] = None,
    ) -> None:
        """
        Write tables to Word document.

        Args:
            output_path: Output file path
            tables: Tables to write
            metadata: Document metadata
        """
        doc = Document()

        # Set page orientation
        self._set_page_orientation(doc)

        # Add document title if available
        if metadata and metadata.title:
            self._add_title(doc, metadata.title)

        # Write each table
        for table_idx, table in enumerate(tables):
            # Add section title if available
            if table.section_title:
                self._add_heading(doc, table.section_title, level=1)

            # Add table
            self._write_table(doc, table)

            # Add page break if requested (except after last table)
            if self.page_break_per_table and table_idx < len(tables) - 1:
                doc.add_page_break()

        # Save document
        doc.save(output_path)

    def _set_page_orientation(self, doc: Document) -> None:
        """Set page orientation for document."""
        section = doc.sections[0]

        if self.orientation == WordOrientation.LANDSCAPE:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            # Swap width and height for landscape
            section.page_width, section.page_height = (
                section.page_height,
                section.page_width,
            )
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT

    def _add_title(self, doc: Document, title: str) -> None:
        """Add document title."""
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply theme color if available
        if self.theme:
            for run in title_para.runs:
                run.font.color.rgb = self._hex_to_rgb(self.theme.palette.primary)

    def _add_heading(self, doc: Document, text: str, level: int = 1) -> None:
        """Add a heading."""
        heading = doc.add_heading(text, level=level)

        # Apply theme color if available
        if self.theme:
            for run in heading.runs:
                run.font.color.rgb = self._hex_to_rgb(self.theme.palette.primary)

    def _write_table(self, doc: Document, table: ExtractedTable) -> None:
        """
        Write a table to the document.

        Args:
            doc: Document
            table: Table to write
        """
        if table.is_empty:
            return

        # Create Word table
        num_rows = len(table.data)
        num_cols = table.schema.column_count

        word_table = doc.add_table(rows=num_rows, cols=num_cols)
        word_table.style = "Light Grid Accent 1"

        # Fill table data
        for row_idx, row_data in enumerate(table.data):
            for col_idx, cell_value in enumerate(row_data):
                if col_idx < num_cols:
                    cell = word_table.cell(row_idx, col_idx)
                    cell.text = str(cell_value) if cell_value else ""

        # Style header row
        if table.schema.has_header:
            self._style_header_row(word_table)

        # Add spacing after table
        doc.add_paragraph()

    def _style_header_row(self, table: DocxTable) -> None:
        """Apply styling to header row."""
        if not table.rows:
            return

        header_row = table.rows[0]

        # Get theme colors or use defaults
        if self.theme:
            bg_color = self._hex_to_rgb(self.theme.palette.primary)
        else:
            bg_color = RGBColor(11, 31, 59)  # Corporate navy blue

        for cell in header_row.cells:
            # Set background color
            self._set_cell_background(cell, bg_color)

            # Set text color and bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White
                    run.font.size = Pt(11)

                # Center align
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _set_cell_background(self, cell, color: RGBColor) -> None:
        """Set cell background color."""
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), f"{color.red:02x}{color.green:02x}{color.blue:02x}")
        cell._element.get_or_add_tcPr().append(shading_elm)

    @staticmethod
    def _hex_to_rgb(hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor."""
        hex_color = hex_color.lstrip("#")
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
