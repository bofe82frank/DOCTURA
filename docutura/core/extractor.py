"""
Core extraction engine for DocTura Desktop.

Handles PDF, DOCX, and image file ingestion with OCR support.
"""

import io
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pdfplumber
from PIL import Image
from pypdf import PdfReader

try:
    import pytesseract

    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False

from docutura.core.models import ExtractedTable, SegmentationStrategy, TableSchema


class DocumentExtractor:
    """Main extraction engine for documents."""

    def __init__(self, enable_ocr: bool = False, ocr_language: str = "eng"):
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language

        if enable_ocr and not PYTESSERACT_AVAILABLE:
            raise ImportError(
                "pytesseract is required for OCR. Install it with: pip install pytesseract"
            )

    def detect_file_type(self, file_path: Path) -> str:
        """
        Detect file type from extension.

        Args:
            file_path: Path to file

        Returns:
            File type: 'pdf', 'docx', 'image', 'audio', or 'unknown'
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return "pdf"
        elif suffix in [".docx", ".doc"]:
            return "docx"
        elif suffix in [".png", ".jpg", ".jpeg", ".tiff", ".tif"]:
            return "image"
        elif suffix in [".wav", ".mp3", ".m4a"]:
            return "audio"
        else:
            return "unknown"

    def needs_ocr(self, file_path: Path) -> bool:
        """
        Determine if file needs OCR.

        Args:
            file_path: Path to file

        Returns:
            True if OCR is needed
        """
        file_type = self.detect_file_type(file_path)

        # Images always need OCR
        if file_type == "image":
            return True

        # Check if PDF has extractable text
        if file_type == "pdf":
            try:
                reader = PdfReader(str(file_path))
                if len(reader.pages) > 0:
                    text = reader.pages[0].extract_text()
                    return len(text.strip()) < 50  # Arbitrary threshold
            except Exception:
                return True

        return False

    def extract_from_pdf(self, file_path: Path) -> Tuple[List[Dict[str, Any]], List[str]]:
        """
        Extract tables and text from PDF.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (tables_data, page_texts)
        """
        tables_data = []
        page_texts = []

        with pdfplumber.open(str(file_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract text
                text = page.extract_text() or ""
                page_texts.append(text)

                # Extract tables using pdfplumber
                tables = page.extract_tables()

                for table_idx, table in enumerate(tables):
                    if table and len(table) > 0:
                        # Clean table data
                        cleaned_table = self._clean_table_data(table)

                        if cleaned_table:
                            tables_data.append(
                                {
                                    "data": cleaned_table,
                                    "page": page_num,
                                    "table_index": table_idx,
                                    "source": "pdfplumber",
                                }
                            )

        return tables_data, page_texts

    def extract_from_pdf_with_ocr(
        self, file_path: Path
    ) -> Tuple[List[Dict[str, Any]], List[str]]:
        """
        Extract from PDF using OCR.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (tables_data, page_texts)
        """
        if not PYTESSERACT_AVAILABLE:
            raise ImportError("pytesseract is required for OCR")

        from pdf2image import convert_from_path

        tables_data = []
        page_texts = []

        # Convert PDF to images
        images = convert_from_path(str(file_path))

        for page_num, image in enumerate(images, start=1):
            # Perform OCR
            text = pytesseract.image_to_string(image, lang=self.ocr_language)
            page_texts.append(text)

            # For now, we don't extract tables from OCR
            # This would require more sophisticated table detection
            # Future enhancement: use table detection ML models

        return tables_data, page_texts

    def extract_from_image(self, file_path: Path) -> str:
        """
        Extract text from image using OCR.

        Args:
            file_path: Path to image file

        Returns:
            Extracted text
        """
        if not PYTESSERACT_AVAILABLE:
            raise ImportError("pytesseract is required for OCR")

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang=self.ocr_language)
        return text

    def extract_from_docx(self, file_path: Path) -> Tuple[List[Dict[str, Any]], str]:
        """
        Extract tables and text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (tables_data, text)
        """
        from docx import Document

        doc = Document(file_path)
        tables_data = []
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []

            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                cleaned_table = self._clean_table_data(table_data)
                if cleaned_table:
                    tables_data.append(
                        {
                            "data": cleaned_table,
                            "page": 1,  # DOCX doesn't have page numbers in the same way
                            "table_index": table_idx,
                            "source": "python-docx",
                        }
                    )

        full_text = "\n".join(text_parts)
        return tables_data, full_text

    def _clean_table_data(self, table: List[List[Any]]) -> List[List[Any]]:
        """
        Clean and normalize table data.

        Args:
            table: Raw table data

        Returns:
            Cleaned table data
        """
        if not table:
            return []

        cleaned = []

        for row in table:
            if not row:
                continue

            # Clean cell values
            cleaned_row = []
            for cell in row:
                if cell is None:
                    cleaned_row.append("")
                elif isinstance(cell, str):
                    cleaned_row.append(cell.strip())
                else:
                    cleaned_row.append(str(cell).strip())

            # Skip completely empty rows
            if any(cell for cell in cleaned_row):
                cleaned.append(cleaned_row)

        return cleaned

    def create_page_preserved_tables(
        self, tables_data: List[Dict[str, Any]]
    ) -> List[ExtractedTable]:
        """
        Create page-preserved table representations.

        Args:
            tables_data: Raw extracted tables with page information

        Returns:
            List of page-preserved ExtractedTable objects
        """
        page_tables_dict: Dict[int, List[List[Any]]] = {}

        # Group tables by page
        for table_dict in tables_data:
            page = table_dict["page"]
            data = table_dict["data"]

            if page not in page_tables_dict:
                page_tables_dict[page] = []

            # Merge multiple tables from same page
            if page_tables_dict[page]:
                # Add a blank row separator
                page_tables_dict[page].append([""] * len(data[0]))

            page_tables_dict[page].extend(data)

        # Convert to ExtractedTable objects
        page_tables = []
        for page_num in sorted(page_tables_dict.keys()):
            data = page_tables_dict[page_num]

            if data:
                # Try to detect headers (first row with all non-empty cells)
                has_header = False
                headers = []

                if data and all(cell.strip() for cell in data[0]):
                    has_header = True
                    headers = data[0]

                schema = TableSchema(
                    headers=headers if headers else [f"Column_{i+1}" for i in range(len(data[0]))],
                    column_count=len(data[0]) if data else 0,
                    has_header=has_header,
                    header_row_indices=[0] if has_header else [],
                )

                table = ExtractedTable(
                    data=data,
                    schema=schema,
                    source_pages=[page_num],
                    table_type="page_preserved",
                )

                page_tables.append(table)

        return page_tables

    def extract(
        self, file_path: Path
    ) -> Tuple[List[Dict[str, Any]], List[str], Dict[str, Any]]:
        """
        Main extraction method.

        Args:
            file_path: Path to file

        Returns:
            Tuple of (tables_data, page_texts, context)
        """
        file_type = self.detect_file_type(file_path)
        context = {"file_type": file_type, "file_path": str(file_path)}

        if file_type == "pdf":
            needs_ocr = self.enable_ocr and self.needs_ocr(file_path)

            if needs_ocr:
                tables_data, page_texts = self.extract_from_pdf_with_ocr(file_path)
                context["extraction_method"] = "ocr"
            else:
                tables_data, page_texts = self.extract_from_pdf(file_path)
                context["extraction_method"] = "native"

        elif file_type == "docx":
            tables_data, full_text = self.extract_from_docx(file_path)
            page_texts = [full_text]
            context["extraction_method"] = "python-docx"

        elif file_type == "image":
            if not self.enable_ocr:
                raise ValueError("OCR must be enabled to process image files")

            text = self.extract_from_image(file_path)
            tables_data = []
            page_texts = [text]
            context["extraction_method"] = "ocr"

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        context["tables_extracted"] = len(tables_data)
        context["pages_processed"] = len(page_texts)

        return tables_data, page_texts, context
